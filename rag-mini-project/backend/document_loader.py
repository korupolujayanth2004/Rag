# backend/document_loader.py
import fitz  # PyMuPDF for PDF
from docx import Document as DocxDocument # Aliased to avoid name conflict with our Document class
import openpyxl
import csv
import json
from bs4 import BeautifulSoup
from pydantic import BaseModel, Field
from typing import Dict, Any, List
from pathlib import Path
import uuid # Essential for generating unique IDs

# Import your chunker utility
from backend.chunker import chunk_text

# --- Define the Document class ---
# This Pydantic model defines the structure for each processed document chunk.
class Document(BaseModel):
    text: str # The content of the chunk
    metadata: Dict[str, Any] = Field(default_factory=dict) # Metadata like source, page number
    chunk_id: str = Field(default_factory=lambda: str(uuid.uuid4())) # Unique ID for this specific chunk

# --- Document Loading and Chunking Functions ---
def extract_text(file_path: Path, content: bytes) -> List[Document]:
    """
    Extracts raw text from various document types, then processes this text
    into smaller, manageable chunks using the 'chunk_text' utility.
    
    Args:
        file_path (Path): The path object for the uploaded file (used for name/extension).
        content (bytes): The raw byte content of the uploaded file.

    Returns:
        List[Document]: A list of Document objects, each representing a text chunk.
    """
    raw_texts_with_metadata = [] # Temporarily stores extracted text before final chunking
    file_type = file_path.suffix.lower().lstrip(".")
    filename = file_path.name

    try:
        # --- PDF Handling ---
        if file_type == "pdf":
            pdf_document = fitz.open(stream=content, filetype="pdf")
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    raw_texts_with_metadata.append(
                        {
                            "text": text,
                            "metadata": {
                                "source": filename,
                                "page_number": page_num + 1,
                                "file_type": "pdf"
                            }
                        }
                    )
            pdf_document.close()

        # --- Text File Handling ---
        elif file_type == "txt":
            text = content.decode('utf-8')
            if text.strip():
                raw_texts_with_metadata.append(
                    {
                        "text": text,
                        "metadata": {
                            "source": filename,
                            "file_type": "txt"
                        }
                    }
                )

        # --- DOCX (Word) Handling ---
        elif file_type == "docx":
            from io import BytesIO
            doc = DocxDocument(BytesIO(content))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text = "\n".join(full_text)
            if text.strip():
                raw_texts_with_metadata.append(
                    {
                        "text": text,
                        "metadata": {
                            "source": filename,
                            "file_type": "docx"
                        }
                    }
                )

        # --- XLSX (Excel) Handling ---
        elif file_type == "xlsx":
            from io import BytesIO
            workbook = openpyxl.load_workbook(BytesIO(content))
            all_sheets_text = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                for row in sheet.iter_rows():
                    row_values = [str(cell.value) if cell.value is not None else "" for cell in row]
                    sheet_text.append("\t".join(row_values))
                all_sheets_text.append(f"Sheet: {sheet_name}\n" + "\n".join(sheet_text))
            text = "\n\n".join(all_sheets_text)
            if text.strip():
                raw_texts_with_metadata.append(
                    {
                        "text": text,
                        "metadata": {
                            "source": filename,
                            "file_type": "xlsx"
                        }
                    }
                )

        # --- CSV Handling ---
        elif file_type == "csv":
            from io import StringIO
            decoded_content = content.decode('utf-8')
            reader = csv.reader(StringIO(decoded_content))
            csv_data = [",".join(row) for row in reader]
            text = "\n".join(csv_data)
            if text.strip():
                raw_texts_with_metadata.append(
                    {
                        "text": text,
                        "metadata": {
                            "source": filename,
                            "file_type": "csv"
                        }
                    }
                )

        # --- JSON Handling ---
        elif file_type == "json":
            decoded_content = content.decode('utf-8')
            json_data = json.loads(decoded_content)
            text = json.dumps(json_data, indent=2) # Pretty-print JSON for readability
            if text.strip():
                raw_texts_with_metadata.append(
                    {
                        "text": text,
                        "metadata": {
                            "source": filename,
                            "file_type": "json"
                        }
                    }
                )

        # --- HTML Handling ---
        elif file_type == "html":
            decoded_content = content.decode('utf-8')
            soup = BeautifulSoup(decoded_content, 'html.parser')
            text = soup.get_text(separator='\n', strip=True) # Extract readable text, remove extra whitespace
            if text.strip():
                raw_texts_with_metadata.append(
                    {
                        "text": text,
                        "metadata": {
                            "source": filename,
                            "file_type": "html"
                        }
                    } # <<< FIXED: Changed ')' to '}' here!
                )

        # --- Fallback for Unsupported Types (attempt to decode as plain text) ---
        else:
            print(f"Unsupported file type: {file_type}. Attempting to decode as plain text.")
            try:
                text = content.decode('utf-8')
                if text.strip():
                    raw_texts_with_metadata.append(
                        {
                            "text": text,
                            "metadata": {
                                "source": filename,
                                "file_type": f"unsupported_{file_type}"
                            }
                        }
                    )
            except UnicodeDecodeError:
                print(f"Could not decode {filename} as UTF-8 text. Skipping.")
                pass # If it cannot be decoded, simply skip this file

    except Exception as e:
        print(f"Error processing file {filename}: {e}")
        # In a production app, you might want to log this error more formally
        # or return an error status for this specific file.

    # --- Apply Chunking to all extracted raw texts ---
    final_documents = []
    for item in raw_texts_with_metadata:
        base_text = item["text"]
        base_metadata = item["metadata"]

        # Use the chunk_text function from backend.chunker to split the raw text
        chunks_from_chonkie = chunk_text(base_text)

        for chunk_content in chunks_from_chonkie:
            if chunk_content.strip(): # Only add non-empty chunks
                # Create a new Document object for each chunk, preserving original metadata
                final_documents.append(
                    Document(
                        text=chunk_content,
                        metadata=base_metadata.copy() # Use .copy() to prevent modifying shared metadata dicts
                    )
                )

    return final_documents